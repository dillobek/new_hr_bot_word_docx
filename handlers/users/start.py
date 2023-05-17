from aiogram import types
from aiogram.dispatcher.filters.builtin import CommandStart
from aiogram.dispatcher import FSMContext
from aiogram.types import ReplyKeyboardRemove
from docx import Document
from docx.shared import Inches, Pt

from keyboards.default.all_def_button import boshla
from keyboards.inline.all_inline_button import edu, family, jobs, rus
from loader import dp, bot
from states.all_states import reg


@dp.message_handler(CommandStart())
async def bot_start(message: types.Message):
    await message.answer(f"Assalamu alaykum <b>BILLI</b> bolalar akademiyasi HR botiga hush kelibsiz!\n\nIltimos quyida so'raladigan savollarga aniqlik bilan javob bering!")
    await message.answer(f"Arizani to'ldirishni boshlash uchun <b>Boshlash</b> tugmasini bosing", reply_markup=boshla)

@dp.message_handler(text="📝 Boshlash")
async def start_ask(message:types.Message):
    await message.answer(f"Ismingizni kiriting.\n\nMasalan: <b>Saydullo Xaydarov</b>", reply_markup=ReplyKeyboardRemove())
    await reg.name.set()

@dp.message_handler(state=reg.name)
async def save_name(message:types.Message, state:FSMContext):
    name = message.text
    await state.update_data(
        {"name":name}
    )
    await message.answer(f"Tug'ulgan yilingizni kiriting.\n\nMasalan: <b>13.05.2001</b>")
    await reg.year.set()

@dp.message_handler(state=reg.year)
async def save_year(message:types.Message,state:FSMContext):
    year = message.text
    await state.update_data(
        {"year":year}
    )
    await message.answer(f"Telefon raqamingizni kiriting.\n\nMasalan: <b>+998XXXXXXXXX</b>")
    await reg.phone.set()

@dp.message_handler(state=reg.phone)
async def save_phone(message:types.Message,state:FSMContext):
    phone = message.text
    try:
        if phone.startswith("+998") and len(phone) == 13 and int(message.text):
            await state.update_data(
                {"phone": phone}
            )
            await message.answer("Qo'shimcha raqam kiriting.\n\nMasalan: <b>+998XXXXXXXXX</b>")
            await reg.sphone.set()
        else:
            await message.answer("Iltimos tekshirib qayta kiriting")
    except Exception:
        await message.answer("Iltimos tekshirib qayta kiriting")

@dp.message_handler(state=reg.sphone)
async def save_sphone(message:types.Message,state:FSMContext):
    sphone = message.text
    try:
        if sphone.startswith("+998") and len(sphone) == 13 and int(message.text):
            await state.update_data(
                {"sphone": sphone}
            )
            await message.answer("Ma'lumotingiz qanday ?", reply_markup=edu)
            await reg.graduate.set()
        else:
            await message.answer("Iltimos tekshirib qayta kiriting")
    except Exception:
        await message.answer("Iltimos tekshirib qayta kiriting")

@dp.callback_query_handler(state=reg.graduate)
async def save_graduate(call:types.CallbackQuery,state:FSMContext):
    graduate = call.data
    await state.update_data(
        {"graduate":graduate}
    )
    await call.message.answer(f"Qaysi Ta'lim muassasasini tugatgansiz ?")
    await reg.university.set()

@dp.message_handler(state=reg.university)
async def save_uni(message:types.Message,state:FSMContext):
    university = message.text
    await state.update_data(
        {"university":university}
    )
    await message.answer(f"Oilaviy ahvolingiz", reply_markup=family)
    await reg.sud.set()

@dp.callback_query_handler(state=reg.sud)
async def save_sud(call:types.CallbackQuery,state:FSMContext):
    sud = call.data
    await state.update_data(
        {"sud":sud}
    )
    await call.message.delete()
    await call.message.answer(f"Oldin qanday tashkilotlarda ishlagansiz ?")
    await reg.oldjob.set()

@dp.message_handler(state=reg.oldjob)
async def save_oldjob(message:types.Message, state:FSMContext):
    oldjob = message.text
    await state.update_data(
        {"oldjob":oldjob}
    )
    await message.answer(f"Bizning korxonada qanday yo'nalishda ishlamoqchisiz ?",reply_markup=jobs)
    await reg.newjob.set()

@dp.callback_query_handler(state=reg.newjob)
async def save_newjob(call:types.CallbackQuery,state:FSMContext):
    newjob = call.data
    await state.update_data(
        {"newjob":newjob}
    )
    await call.message.delete()
    await call.message.answer(f"Oldingi maoshingizni kiriting.")
    await reg.oldsalary.set()

@dp.message_handler(state=reg.oldsalary)
async def save_oldsalary(message:types.Message, state:FSMContext):
    oldsalary = message.text
    await state.update_data(
        {"oldsalary":oldsalary}
    )
    await message.answer(f"Yangiz ishingizdan qancha maosh kutyapsiz ?")
    await reg.newsalary.set()

@dp.message_handler(state=reg.newsalary)
async def save_newsalary(message:types.Message, state:FSMContext):
    newsalary = message.text
    await state.update_data(
        {"newsalary":newsalary}
    )

    await message.answer(f"Bizining tashkilotda qancha muddat ishlamoqchisiz ?")
    await reg.period.set()

@dp.message_handler(state=reg.period)
async def save_period(message:types.message,state:FSMContext):
    period = message.text
    await state.update_data(
        {"period":period}
    )
    await message.answer(f"Qaysi tillarni bilasiz ?")
    await reg.lang.set()

@dp.message_handler(state=reg.lang)
async def save_lang(message:types.Message, state:FSMContext):
    lang = message.text
    await state.update_data(
        {"lang":lang}
    )
    await message.answer(f"Ushbu tilni qay darajada bilasiz ?", reply_markup=rus)
    await reg.rate.set()

@dp.callback_query_handler(state=reg.rate)
async def save_rate(call:types.CallbackQuery,state:FSMContext):
    rate = call.data
    await state.update_data(
        {"rate":rate}
    )
    await call.message.answer(f"Bizni qanday topdingiz ?")
    await reg.find.set()

@dp.message_handler(state=reg.find)
async def save_find(message:types.Message,state:FSMContext):
    find = message.text
    await state.update_data(
        {"find":find}
    )
    await message.answer(f"Nima uchun sizni ishga olishimiz kerak ?")
    await reg.why.set()

@dp.message_handler(state=reg.why)
async def save_why(message:types.Message,state:FSMContext):
    why = message.text
    await state.update_data(
        {"why":why}
    )
    await message.answer(f"Arizani yakunlash uchun rasm yuboring")
    await reg.photo.set()

@dp.message_handler(state=reg, content_types=types.ContentType.PHOTO)
async def enter_photo(message:types.Message, state:FSMContext):
    photo = message.photo[-1].file_id
    await state.update_data(
        {"photo":photo}
    )

    data = await state.get_data()

    name = data.get('name')
    year = data.get('year')
    phone = data.get('phone')
    sphone = data.get('sphone')
    graduate = data.get('graduate')
    university = data.get('university')
    family = data.get('family')
    sud = data.get('sud')
    oldjob = data.get('oldjob')
    newjob = data.get('newjob')
    oldsalary = data.get('oldsalary')
    newsalary = data.get('newsalary')
    period = data.get('period')
    lang = data.get('lang')
    rate = data.get('rate')
    find = data.get('find')
    why = data.get('why')

    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times'
    font.size = Pt(13)

    document.add_paragraph("SHAXSIY MA'LUMOTLARI", style='Body Text').alignment = 1
    document.add_paragraph(f"Ism,familiyasi:  {name}", style='Body Text')
    document.add_paragraph(f"Tug'ilgan sanasi:  {year}", style='Body Text')
    document.add_paragraph(f"Telefon raqami:  {phone}", style='Body Text')
    document.add_paragraph(f"Qo'shimcha tel:  {sphone}", style='Body Text')
    document.add_paragraph(f"Ma'lumoti:  {graduate}", style='Body Text')
    document.add_paragraph(f"Tugatgan ta'lim muassasasi:  {university}", style='Body Text')
    document.add_paragraph(f"Oilaviy axvoli:  {family}", style='Body Text')
    document.add_paragraph(f"Sudlanganmi:  {sud}", style='Body Text')
    document.add_paragraph(f"ISH HAQIDA MA'LUMOT: ", style='Body Text').alignment = 1

    table = document.add_table(rows=9, cols=2)
    table.rows[0].cells[0].text = 'Qaysi tashkilotlarda va qaysi lavozimlarda ishlagan:'
    table.rows[1].cells[0].text = 'Bizning tashkilotda qaysi lavozimda ishlamoqchi:'
    table.rows[2].cells[0].text = "Oxirgi ish o'rnidagi oylik maoshi:"
    table.rows[3].cells[0].text = 'Bizda qancha miqdorli maoshga ishlamoqchi:'
    table.rows[4].cells[0].text = 'Bizning tashkilotda qancha muddat ishlamoqchi:'
    table.rows[5].cells[0].text = 'Qaysi tillarni biladi:'
    table.rows[6].cells[0].text = 'Bu tillarni qay darajada biladi:'
    table.rows[7].cells[0].text = 'Biz haqimizda qayerdan eshitgan:'
    table.rows[8].cells[0].text = 'Nima uchun aynan sizni ishga olishimiz kerak? '
    table.style = 'TableGrid'
    table.rows[0].cells[1].text = f"{oldjob}"
    table.rows[1].cells[1].text = f"{newjob}"
    table.rows[2].cells[1].text = f"{oldsalary}"
    table.rows[3].cells[1].text = f"{newsalary}"
    table.rows[4].cells[1].text = f"{period}"
    table.rows[5].cells[1].text = f"{lang}"
    table.rows[6].cells[1].text = f"{rate}"
    table.rows[7].cells[1].text = f"{find}"
    table.rows[8].cells[1].text = f"{why}"

    document.save(f'{message.from_user.id}.docx')
    await bot.send_photo(chat_id=-1001235922480, photo=photo,caption=f"ismi: {name}\n\n")
    await bot.send_document(chat_id=-1001235922480,
                            document=types.InputFile(path_or_bytesio=f"{message.from_user.id}.docx",
                                                     filename="Resume.docx"),
                            caption=f"👤 Ism,familiyasi: {name}\n📞 Telefon raqami:  {phone}\n📝 Bo'lim: {newjob}")
    await state.finish()
    await message.answer(f"Bizning korxonani tanlaganingiz uchun rahmat\n\nBiz siz bilan bog'lanamiz", reply_markup=boshla)












